"""
generate_quote_docx.py -- Native DOCX Quote Generator for ServiceWeb

Reads a structured JSON content file and produces a professionally styled
Word document (.docx) matching the ServiceWeb 3-section quote format.

Usage:
    python scripts/generate_quote_docx.py <content.json> <output.docx>

Dependencies (already installed):
    - python-docx 1.2.0
    - Pillow 12.0.0

JSON Content Schema:
    See the plan file or the SERVICEWEB.md documentation for the full schema.
    The schema is validated against real-world quotes (ADAMS EFT Export).
"""

import json
import sys
import re
import os
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

try:
    from PIL import Image, ImageDraw, ImageFont
    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False

# ---------------------------------------------------------------------------
# Color palette (mapped from HTML design system)
# ---------------------------------------------------------------------------
CLR_PRIMARY = RGBColor(0x66, 0x7E, 0xEA)       # #667eea
CLR_SECONDARY = RGBColor(0x76, 0x4B, 0xA2)     # #764ba2
CLR_DARK = RGBColor(0x34, 0x49, 0x5E)          # #34495e
CLR_BODY = RGBColor(0x33, 0x33, 0x33)          # #333333
CLR_MUTED = RGBColor(0x88, 0x88, 0x88)         # #888888
CLR_BORDER = RGBColor(0xE0, 0xE0, 0xE0)        # #e0e0e0
CLR_LIGHT_BG = RGBColor(0xF8, 0xF9, 0xFA)      # #f8f9fa
CLR_TC_BORDER = RGBColor(0xDE, 0xE2, 0xE6)     # #dee2e6
CLR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CLR_READ = RGBColor(0x2E, 0xCC, 0x71)          # #2ecc71
CLR_WRITE = RGBColor(0xE7, 0x4C, 0x3C)         # #e74c3c
CLR_WARNING_BG = RGBColor(0xFF, 0xF3, 0xCD)    # #fff3cd
CLR_WARNING_BORDER = RGBColor(0xFF, 0xC1, 0x07) # #ffc107
CLR_WARNING_TITLE = RGBColor(0x85, 0x64, 0x04)  # #856404
CLR_INFO_BG = RGBColor(0xE3, 0xF2, 0xFD)       # #e3f2fd
CLR_INFO_BORDER = RGBColor(0x21, 0x96, 0xF3)    # #2196F3
CLR_INFO_TITLE = RGBColor(0x15, 0x65, 0xC0)     # #1565C0
CLR_SUCCESS_BG = RGBColor(0xD4, 0xED, 0xDA)     # #d4edda
CLR_SUCCESS_BORDER = RGBColor(0x28, 0xA7, 0x45)  # #28a745
CLR_SUCCESS_TITLE = RGBColor(0x15, 0x57, 0x24)   # #155724
CLR_BUFFER_BG = RGBColor(0xFF, 0xF8, 0xE1)      # #fff8e1
CLR_BUFFER_TEXT = RGBColor(0x85, 0x64, 0x04)     # #856404
CLR_DISCLAIMER_BG = RGBColor(0xF9, 0xF9, 0xF9)  # #f9f9f9
CLR_SIG_LABEL = RGBColor(0x66, 0x66, 0x66)      # #666666
CLR_HEADING_BORDER = RGBColor(0xE8, 0xE8, 0xE8)  # #e8e8e8
CLR_CODE_BG = RGBColor(0x1E, 0x1E, 0x1E)        # #1e1e1e -- dark code block bg
CLR_CODE_FG = RGBColor(0xD4, 0xD4, 0xD4)        # #d4d4d4 -- code block text
CLR_COMMENT = RGBColor(0x60, 0x8B, 0x4E)        # #608b4e -- code comment

FONT_BODY = "Segoe UI"
FONT_CODE = "Consolas"

# ---------------------------------------------------------------------------
# Lightweight markdown parser
# ---------------------------------------------------------------------------

_MD_PATTERNS = [
    (re.compile(r'\*\*(.+?)\*\*'), 'bold'),
    (re.compile(r'\*(.+?)\*'), 'italic'),
    (re.compile(r'`(.+?)`'), 'code'),
]


def _add_md_runs(paragraph, text, base_color=CLR_BODY, base_size=Pt(11)):
    """Parse lightweight markdown in *text* and add formatted runs to *paragraph*."""
    tokens = _tokenize_md(text)
    for content, fmt in tokens:
        run = paragraph.add_run(content)
        run.font.name = FONT_CODE if fmt == 'code' else FONT_BODY
        run.font.size = base_size
        run.font.color.rgb = base_color
        if fmt == 'bold':
            run.bold = True
        elif fmt == 'italic':
            run.italic = True
        elif fmt == 'code':
            run.font.size = Pt(base_size.pt - 1) if base_size else Pt(10)


def _tokenize_md(text):
    """Return list of (content, format) tuples. format is None, 'bold', 'italic', or 'code'."""
    parts = []
    combined = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    segments = combined.split(text)
    for seg in segments:
        if not seg:
            continue
        if seg.startswith('**') and seg.endswith('**'):
            parts.append((seg[2:-2], 'bold'))
        elif seg.startswith('*') and seg.endswith('*') and not seg.startswith('**'):
            parts.append((seg[1:-1], 'italic'))
        elif seg.startswith('`') and seg.endswith('`'):
            parts.append((seg[1:-1], 'code'))
        else:
            parts.append((seg, None))
    return parts


# ---------------------------------------------------------------------------
# XML / OPC helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, hex_color):
    """Set background shading on a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_paragraph_border_bottom(paragraph, color_hex="E0E0E0", size="6", space="8"):
    """Add a bottom border to a paragraph (used for section separators and h3 headings)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="{space}" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _set_cell_margins(cell, top=0, bottom=0, left=100, right=100):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _set_cell_border(cell, sides=None, color="DEE2E6", size="4"):
    """Set borders on specific sides of a cell."""
    if sides is None:
        sides = ["top", "bottom", "start", "end"]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders_xml = f'<w:tcBorders {nsdecls("w")}>'
    for side in sides:
        borders_xml += f'<w:{side} w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
    borders_xml += '</w:tcBorders>'
    tcPr.append(parse_xml(borders_xml))


def _remove_table_borders(table):
    """Remove all borders from a table (for layout tables)."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


# ---------------------------------------------------------------------------
# Paragraph / text helpers
# ---------------------------------------------------------------------------

def _add_heading_h3(doc, text, color=CLR_PRIMARY):
    """Add an h3-equivalent heading with bottom border."""
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = FONT_BODY
    run.font.color.rgb = color
    _set_paragraph_border_bottom(p, color_hex="E8E8E8", size="4", space="4")
    return p


def _add_heading_h4(doc, text, color=CLR_SECONDARY):
    """Add an h4-equivalent sub-heading."""
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = FONT_BODY
    run.font.color.rgb = color
    return p


def _add_body_paragraph(doc, text, color=CLR_BODY, size=Pt(11)):
    """Add a body text paragraph with markdown parsing."""
    p = doc.add_paragraph()
    p.space_after = Pt(6)
    _add_md_runs(p, text, base_color=color, base_size=size)
    return p


def _add_bullet(doc, text, color=CLR_BODY, size=Pt(11)):
    """Add a bulleted list item with markdown parsing."""
    p = doc.add_paragraph(style='List Bullet')
    p.space_after = Pt(3)
    _add_md_runs(p, text, base_color=color, base_size=size)
    return p


def _add_field_label(doc, text):
    """Add a ServiceWeb field label (small italic gray text)."""
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = FONT_BODY
    run.font.color.rgb = CLR_MUTED
    run.italic = True
    return p


def _add_separator(doc):
    """Add a thin gray horizontal rule between sections."""
    p = doc.add_paragraph()
    p.space_before = Pt(20)
    p.space_after = Pt(6)
    _set_paragraph_border_bottom(p, color_hex="E0E0E0", size="4", space="1")
    return p


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------

def _add_data_table(doc, columns, rows, col_widths=None):
    """Add a styled data table with dark headers."""
    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, col_name in enumerate(columns):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, "34495E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(col_name)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = FONT_BODY
        run.font.color.rgb = CLR_WHITE
        _set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            p = cell.paragraphs[0]
            _add_md_runs(p, str(cell_text), base_color=CLR_BODY, base_size=Pt(10))
            _set_cell_margins(cell, top=40, bottom=40, left=100, right=100)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    return table


# ---------------------------------------------------------------------------
# Section 1: Description
# ---------------------------------------------------------------------------

def _build_section1(doc, data):
    _add_field_label(doc, "Description")

    sec1 = data.get("section1_description", {})

    for para_text in sec1.get("executive_summary", []):
        _add_body_paragraph(doc, para_text)

    screenshots = sec1.get("screenshots", [])
    if screenshots:
        for ss in screenshots:
            _build_screenshot(doc, ss)

    assumptions = sec1.get("assumptions", [])
    if assumptions:
        _add_heading_h3(doc, "Assumptions")
        for item in assumptions:
            _add_bullet(doc, item)

    prereqs = sec1.get("prerequisites", [])
    if prereqs:
        _add_heading_h3(doc, "Prerequisites")
        for item in prereqs:
            _add_bullet(doc, item)


# ---------------------------------------------------------------------------
# Section 2: Specifications (Customer Facing)
# ---------------------------------------------------------------------------

def _build_section2(doc, data):
    _add_field_label(doc, "Specifications")

    sec2 = data.get("section2_specifications", {})

    # Scope of Work
    scope_items = sec2.get("scope_items", [])
    if scope_items:
        _add_heading_h3(doc, "Scope of Work")
        for item in scope_items:
            num = item.get("number", "")
            title = item.get("title", "")
            _add_heading_h4(doc, f"{num}. {title}")
            desc = item.get("description", "")
            if desc:
                _add_body_paragraph(doc, desc)
            for sub in item.get("sub_items", []):
                _add_bullet(doc, sub)

    # Deliverables
    deliverables = sec2.get("deliverables", [])
    if deliverables:
        _add_heading_h3(doc, "Deliverables")
        columns = ["#", "Deliverable", "Type", "Description"]
        rows = []
        for d in deliverables:
            rows.append([
                d.get("id", ""),
                f"**{d.get('deliverable', '')}**",
                d.get("type", ""),
                d.get("description", ""),
            ])
        _add_data_table(doc, columns, rows,
                        col_widths=[Inches(0.4), Inches(1.5), Inches(1.0), Inches(3.6)])

    # Testing Plan
    test_cases = sec2.get("test_cases", [])
    if test_cases:
        _add_heading_h3(doc, "Testing Plan")
        for tc in test_cases:
            _build_test_case(doc, tc)


def _build_test_case(doc, tc):
    """Render a single test case in a bordered box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "F8F9FA")
    _set_cell_border(cell, color="DEE2E6", size="4")
    _set_cell_margins(cell, top=100, bottom=100, left=140, right=140)

    # Title
    p_title = cell.paragraphs[0]
    run_title = p_title.add_run(f"{tc.get('id', '')}: {tc.get('title', '')}")
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.name = FONT_BODY
    run_title.font.color.rgb = CLR_SECONDARY
    p_title.space_after = Pt(6)

    # Optional reference note
    ref = tc.get("reference")
    if ref:
        p_ref = cell.add_paragraph()
        p_ref.space_after = Pt(4)
        run_ref = p_ref.add_run(f"Reference: {ref}")
        run_ref.font.size = Pt(9)
        run_ref.font.name = FONT_BODY
        run_ref.font.color.rgb = CLR_MUTED
        run_ref.italic = True

    # Objective
    _tc_label(cell, "Objective:")
    obj = tc.get("objective", "")
    p_obj = cell.add_paragraph(style='List Bullet')
    p_obj.space_after = Pt(2)
    _add_md_runs(p_obj, obj, base_color=CLR_BODY, base_size=Pt(10))

    # Steps
    _tc_label(cell, "Steps:")
    for step in tc.get("steps", []):
        p_step = cell.add_paragraph(style='List Bullet')
        p_step.space_after = Pt(2)
        _add_md_runs(p_step, step, base_color=CLR_BODY, base_size=Pt(10))

    # Expected Results (supports sub-headed groups)
    expected = tc.get("expected_results", [])
    if expected:
        if isinstance(expected[0], dict):
            for group in expected:
                heading = group.get("heading")
                label_text = f"Expected Results \u2014 {heading}:" if heading else "Expected Results:"
                _tc_label(cell, label_text)
                for item in group.get("items", []):
                    p_er = cell.add_paragraph(style='List Bullet')
                    p_er.space_after = Pt(2)
                    _add_md_runs(p_er, item, base_color=CLR_BODY, base_size=Pt(10))
        else:
            _tc_label(cell, "Expected Results:")
            for item in expected:
                p_er = cell.add_paragraph(style='List Bullet')
                p_er.space_after = Pt(2)
                _add_md_runs(p_er, item, base_color=CLR_BODY, base_size=Pt(10))


def _tc_label(cell, text):
    """Add a test-case section label (Objective:, Steps:, Expected Results:)."""
    p = cell.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = FONT_BODY
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------------------
# Section 3: Technical Specifications
# ---------------------------------------------------------------------------

def _build_section3(doc, data):
    _add_field_label(doc, "Technical Specifications")

    sec3 = data.get("section3_technical", {})

    # Content blocks (ordered sequence)
    for block in sec3.get("content_blocks", []):
        btype = block.get("type", "")
        if btype == "heading":
            level = block.get("level", 3)
            text = block.get("text", "")
            if level == 3:
                _add_heading_h3(doc, text)
            else:
                _add_heading_h4(doc, text)

        elif btype == "sub_heading":
            num = block.get("number", "")
            title = block.get("title", "")
            _add_heading_h4(doc, f"{num}. {title}")
            for para in block.get("paragraphs", []):
                _add_body_paragraph(doc, para)
            for bullet in block.get("bullets", []):
                _add_bullet(doc, bullet)

        elif btype == "table":
            cols = block.get("columns", [])
            rows = block.get("rows", [])
            if cols and rows:
                _add_data_table(doc, cols, rows)

        elif btype == "code_block":
            _build_code_block(doc, block)

        elif btype == "callout":
            _build_callout(doc, block)

    # Architecture diagram
    arch = sec3.get("architecture")
    if arch:
        _add_heading_h3(doc, arch.get("title", "Architecture Overview"))
        if HAS_PILLOW:
            png_bytes = _render_architecture_diagram(arch)
            if png_bytes:
                doc.add_picture(BytesIO(png_bytes), width=Inches(6.5))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            _add_body_paragraph(doc, "[Architecture diagram requires Pillow library]",
                                color=CLR_WARNING_TITLE)

    # Integration Points
    int_points = sec3.get("integration_points", [])
    if int_points:
        _add_heading_h3(doc, "Integration Points")
        columns = ["GSS Module", "Direction", "Integration Points"]
        rows = []
        for ip in int_points:
            direction_text = ip.get("direction", "")
            rows.append([
                f"**{ip.get('module', '')}**",
                direction_text,
                ip.get("description", ""),
            ])
        table = _add_data_table(doc, columns, rows,
                                col_widths=[Inches(1.4), Inches(0.9), Inches(4.2)])
        _colorize_direction_cells(table, col_index=1)

    # Screenshots
    screenshots = sec3.get("screenshots", [])
    if screenshots:
        _add_heading_h3(doc, "Mockup Screenshots")
        for ss in screenshots:
            _build_screenshot(doc, ss)


def _build_code_block(doc, block):
    """Render a dark-background code block."""
    title = block.get("title")
    if title:
        p_title = doc.add_paragraph()
        p_title.space_before = Pt(8)
        p_title.space_after = Pt(2)
        run = p_title.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = FONT_BODY
        run.font.color.rgb = CLR_SECONDARY

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "1E1E1E")
    _set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    _set_cell_border(cell, color="333333", size="4")

    first = True
    for line in block.get("lines", []):
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.space_after = Pt(1)
        p.space_before = Pt(1)

        if "comment" in line:
            run = p.add_run(line["comment"])
            run.font.name = FONT_CODE
            run.font.size = Pt(9)
            run.font.color.rgb = CLR_COMMENT
        elif "text" in line:
            run = p.add_run(line["text"])
            run.font.name = FONT_CODE
            run.font.size = Pt(9)
            run.font.color.rgb = CLR_CODE_FG


def _build_callout(doc, block):
    """Render a callout box (info, warning, or success)."""
    style = block.get("style", "info")
    if style == "warning":
        bg_hex, border_clr, title_clr = "FFF3CD", CLR_WARNING_BORDER, CLR_WARNING_TITLE
    elif style == "success":
        bg_hex, border_clr, title_clr = "D4EDDA", CLR_SUCCESS_BORDER, CLR_SUCCESS_TITLE
    else:
        bg_hex, border_clr, title_clr = "E3F2FD", CLR_INFO_BORDER, CLR_INFO_TITLE

    border_hex = f"{border_clr.red:02X}{border_clr.green:02X}{border_clr.blue:02X}" if hasattr(border_clr, 'red') else "2196F3"

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, bg_hex)
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=140)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders_xml = (
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{border_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_hex}"/>'
        f'  <w:end w:val="single" w:sz="4" w:space="0" w:color="{border_hex}"/>'
        f'  <w:start w:val="single" w:sz="12" w:space="0" w:color="{border_hex}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(parse_xml(borders_xml))

    p_title = cell.paragraphs[0]
    run = p_title.add_run(block.get("title", ""))
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = FONT_BODY
    run.font.color.rgb = title_clr
    p_title.space_after = Pt(4)

    content = block.get("content", "")
    if content:
        p_content = cell.add_paragraph()
        _add_md_runs(p_content, content, base_color=CLR_BODY, base_size=Pt(10))


def _colorize_direction_cells(table, col_index):
    """Colorize READ/WRITE direction badges in an integration points table."""
    for row_idx in range(1, len(table.rows)):
        cell = table.rows[row_idx].cells[col_index]
        text = cell.paragraphs[0].text.strip()
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]

        if "/" in text:
            parts = [x.strip() for x in text.split("/")]
            for i, part in enumerate(parts):
                clr = CLR_READ if "READ" in part.upper() else CLR_WRITE
                run = p.add_run(part)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_BODY
                run.font.color.rgb = clr
                if i < len(parts) - 1:
                    sep = p.add_run(" / ")
                    sep.font.size = Pt(10)
                    sep.font.name = FONT_BODY
                    sep.font.color.rgb = CLR_BODY
        else:
            clr = CLR_READ if "READ" in text.upper() else CLR_WRITE
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = FONT_BODY
            run.font.color.rgb = clr


def _build_screenshot(doc, ss):
    """Embed a screenshot with caption, image, description, and interaction bullets."""
    caption = ss.get("caption", "Screenshot")
    img_path = ss.get("path", "")
    description = ss.get("description", "")
    bullets = ss.get("interactions", [])

    _add_heading_h4(doc, caption)

    if description:
        _add_body_paragraph(doc, description)

    if img_path and os.path.isfile(img_path):
        doc.add_picture(img_path, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p_warn = doc.add_paragraph()
        p_warn.space_after = Pt(4)
        run = p_warn.add_run(f"[Screenshot not found: {img_path}]")
        run.font.size = Pt(10)
        run.font.name = FONT_BODY
        run.font.color.rgb = CLR_WARNING_TITLE
        run.italic = True

    if bullets:
        p_label = doc.add_paragraph()
        p_label.space_before = Pt(6)
        p_label.space_after = Pt(2)
        run = p_label.add_run("User Interactions:")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = FONT_BODY
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        for bullet_text in bullets:
            _add_bullet(doc, bullet_text, size=Pt(10))


# ---------------------------------------------------------------------------
# Quote Details
# ---------------------------------------------------------------------------

def _build_quote_details(doc, data):
    quote = data.get("quote_details", {})
    line_items = quote.get("line_items", [])
    if not line_items:
        return

    p_heading = doc.add_paragraph()
    p_heading.space_before = Pt(16)
    p_heading.space_after = Pt(6)
    run = p_heading.add_run("Quote Details")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = FONT_BODY
    run.font.color.rgb = CLR_BODY

    columns = ["Sequence", "Type", "Hours", "Rate", "Total", "Description"]
    num_rows = len(line_items) + 1  # +1 for total row
    table = doc.add_table(rows=1 + num_rows, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        _set_cell_shading(header_cells[i], "F0F0F0")
        p = header_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col)
        run.font.size = Pt(10)
        run.font.name = FONT_BODY
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.italic = True
        _set_cell_margins(header_cells[i], top=40, bottom=40, left=60, right=60)

    # Line item rows
    for r_idx, item in enumerate(line_items):
        row = table.rows[r_idx + 1]
        is_buffer = item.get("is_buffer", False)
        values = [
            str(item.get("sequence", "")) if not is_buffer else "",
            item.get("type", ""),
            str(item.get("hours", "")),
            f"${item.get('rate', 0):,.2f}",
            f"${item.get('total', 0):,.2f}",
            item.get("description", ""),
        ]
        for c_idx, val in enumerate(values):
            cell = row.cells[c_idx]
            if is_buffer:
                _set_cell_shading(cell, "FFF8E1")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 5 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = FONT_BODY
            run.font.color.rgb = CLR_BUFFER_TEXT if is_buffer else CLR_BODY
            if is_buffer:
                run.italic = True
            _set_cell_margins(cell, top=30, bottom=30, left=60, right=60)

    # Total row
    total_row = table.rows[-1]
    total_vals = [
        "", "Total",
        str(quote.get("total_hours", "")),
        "",
        f"${quote.get('total_amount', 0):,.2f}",
        "",
    ]
    for c_idx, val in enumerate(total_vals):
        cell = total_row.cells[c_idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = FONT_BODY
        run.font.color.rgb = CLR_BODY
        _set_cell_margins(cell, top=40, bottom=40, left=60, right=60)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        top_border = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="8" w:space="0" w:color="333333"/>'
            '</w:tcBorders>'
        )
        tcPr.append(top_border)


# ---------------------------------------------------------------------------
# Disclaimers
# ---------------------------------------------------------------------------

STANDARD_DISCLAIMERS = [
    ("GSS Version Requirement:", "GSS version 2020.1 or higher required. Compatibility with earlier versions is not guaranteed and may require additional development effort."),
    ("Purchase Order and Billing:", "A purchase order is required in full before commencement of work. For multi-month projects, billing occurs monthly for hours worked. Additional purchase orders will be requested for any approved overages beyond the authorized budget."),
    ("Acceptance of Project Scope/Quotes and Change Requests:", "Acceptance of this quote confirms that the customer has reviewed and agrees to the scope of work described herein. Any changes to scope, requirements, or deliverables after acceptance are subject to the change request and requote process."),
    ("Project Cancellation:", "The customer may cancel the project at any time. All costs incurred to date of cancellation will be billed. Unused hours are not carried forward."),
    ("Training:", "Training is not included in this quote unless explicitly stated as a line item. Training is available as a separate engagement or virtual session at standard rates."),
    ("Terms and Conditions/Service:", "This project will be serviced per the existing Software Licensing Agreement between the customer and Global Shop Solutions. Custom scripts and configurations developed under this engagement are the property of Global Shop Solutions."),
]


def _build_disclaimers(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "F9F9F9")
    _set_cell_border(cell, color="E0E0E0", size="4")
    _set_cell_margins(cell, top=120, bottom=120, left=160, right=160)

    p_title = cell.paragraphs[0]
    p_title.space_after = Pt(8)
    run = p_title.add_run("Standard Terms & Conditions")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = FONT_BODY
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for i, (label, text) in enumerate(STANDARD_DISCLAIMERS, 1):
        p = cell.add_paragraph()
        p.space_after = Pt(3)
        run_num = p.add_run(f"{i}. ")
        run_num.font.size = Pt(9)
        run_num.font.name = FONT_BODY
        run_num.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_label = p.add_run(label + " ")
        run_label.bold = True
        run_label.font.size = Pt(9)
        run_label.font.name = FONT_BODY
        run_label.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run_text = p.add_run(text)
        run_text.font.size = Pt(9)
        run_text.font.name = FONT_BODY
        run_text.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ---------------------------------------------------------------------------
# Signature
# ---------------------------------------------------------------------------

def _build_signature(doc):
    p_intro = doc.add_paragraph()
    p_intro.space_before = Pt(20)
    p_intro.space_after = Pt(12)
    _set_paragraph_border_bottom(p_intro, color_hex="DDDDDD", size="2", space="12")
    run = p_intro.add_run(
        "By signing below, both parties agree to the scope, terms, and authorized budget defined in this quote."
    )
    run.font.size = Pt(10)
    run.font.name = FONT_BODY
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p_line = doc.add_paragraph()
    p_line.space_before = Pt(30)
    p_line.space_after = Pt(2)
    run = p_line.add_run("_" * 50)
    run.font.size = Pt(11)
    run.font.name = FONT_BODY
    run.font.color.rgb = CLR_BODY

    p_label = doc.add_paragraph()
    p_label.space_after = Pt(0)
    run = p_label.add_run("Authorized Signature          Date")
    run.font.size = Pt(9)
    run.font.name = FONT_BODY
    run.font.color.rgb = CLR_SIG_LABEL


# ---------------------------------------------------------------------------
# Architecture Diagram Renderer (Pillow)
# ---------------------------------------------------------------------------

def _get_font(size, bold=False):
    """Load Segoe UI TrueType font, falling back to default if unavailable."""
    font_paths = [
        r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\segoeui.ttf",
    ]
    for fp in font_paths:
        if os.path.isfile(fp):
            try:
                return ImageFont.truetype(fp, size)
            except Exception:
                pass
    return ImageFont.load_default()


def _draw_rounded_rect(draw, xy, radius, fill, outline=None):
    """Draw a rounded rectangle with optional gradient fill."""
    x0, y0, x1, y1 = xy
    if isinstance(fill, tuple) and len(fill) == 2:
        base, end = fill
        height = y1 - y0
        for i in range(height):
            ratio = i / max(height - 1, 1)
            r = int(base[0] + (end[0] - base[0]) * ratio)
            g = int(base[1] + (end[1] - base[1]) * ratio)
            b = int(base[2] + (end[2] - base[2]) * ratio)
            draw.line([(x0 + radius, y0 + i), (x1 - radius, y0 + i)], fill=(r, g, b))
        draw.rounded_rectangle(xy, radius=radius, fill=None, outline=outline, width=2 if outline else 0)
        inner = (x0 + 2, y0 + 2, x1 - 2, y1 - 2)
        mask = Image.new("L", (x1 - x0, y1 - y0), 0)
        mask_draw = ImageDraw.Draw(mask)
        mask_draw.rounded_rectangle((0, 0, x1 - x0, y1 - y0), radius=radius, fill=255)
        return
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=2 if outline else 0)


BOX_PRESETS = {
    "purple": {"fill": ((0x66, 0x7E, 0xEA), (0x76, 0x4B, 0xA2)), "text": (255, 255, 255)},
    "dark": {"fill": (0x2C, 0x3E, 0x50), "text": (255, 255, 255)},
    "amber": {"fill": (0xFF, 0xF3, 0xCD), "text": (0x85, 0x64, 0x04), "outline": (0xFF, 0xC1, 0x07)},
    "green": {"fill": (0xD4, 0xED, 0xDA), "text": (0x15, 0x57, 0x24), "outline": (0x28, 0xA7, 0x45)},
}


def _render_architecture_diagram(arch_data):
    """Render architecture diagram as PNG bytes using Pillow."""
    if not HAS_PILLOW:
        return None

    boxes = arch_data.get("boxes", [])
    connections = arch_data.get("connections", [])
    subtitle = arch_data.get("subtitle", "")

    if not boxes:
        return None

    rows_map = {}
    for box in boxes:
        r = box.get("row", 0)
        rows_map.setdefault(r, []).append(box)
    sorted_rows = sorted(rows_map.keys())

    BOX_W, BOX_H = 240, 100
    GAP_X, GAP_Y = 100, 60
    ARROW_W = 140
    MARGIN = 40

    max_cols = max(len(rows_map[r]) for r in sorted_rows)
    first_row_count = len(rows_map.get(sorted_rows[0], []))
    first_row_w = first_row_count * BOX_W + (first_row_count - 1) * (GAP_X + ARROW_W) if first_row_count > 1 else BOX_W

    img_w = max(first_row_w + MARGIN * 2, max_cols * (BOX_W + GAP_X) + MARGIN * 2)
    img_h = MARGIN + 30 + len(sorted_rows) * (BOX_H + GAP_Y) + MARGIN

    img = Image.new("RGB", (img_w, img_h), (0xF8, 0xF9, 0xFA))
    draw = ImageDraw.Draw(img)

    font_title = _get_font(13, bold=True)
    font_name = _get_font(14, bold=True)
    font_sub = _get_font(11)
    font_detail = _get_font(10)
    font_arrow = _get_font(11, bold=True)
    font_proto = _get_font(9)

    if subtitle:
        tw = draw.textlength(subtitle, font=font_sub)
        draw.text(((img_w - tw) / 2, MARGIN - 10), subtitle, fill=(0x66, 0x66, 0x66), font=font_sub)

    y_offset = MARGIN + 25
    row_box_positions = {}

    for row_idx in sorted_rows:
        row_boxes = rows_map[row_idx]
        if row_idx == sorted_rows[0] and len(row_boxes) >= 2:
            total_w = len(row_boxes) * BOX_W + (len(row_boxes) - 1) * (GAP_X + ARROW_W)
            start_x = (img_w - total_w) / 2
        else:
            total_w = len(row_boxes) * BOX_W + (len(row_boxes) - 1) * GAP_X
            start_x = (img_w - total_w) / 2

        x = start_x
        positions = []
        for i, box in enumerate(row_boxes):
            preset = BOX_PRESETS.get(box.get("color", "purple"), BOX_PRESETS["purple"])
            fill = preset["fill"]
            text_clr = preset["text"]
            outline = preset.get("outline")

            xy = (x, y_offset, x + BOX_W, y_offset + BOX_H)

            if isinstance(fill, tuple) and len(fill) == 2 and isinstance(fill[0], tuple):
                for dy in range(BOX_H):
                    ratio = dy / max(BOX_H - 1, 1)
                    r = int(fill[0][0] + (fill[1][0] - fill[0][0]) * ratio)
                    g = int(fill[0][1] + (fill[1][1] - fill[0][1]) * ratio)
                    b = int(fill[0][2] + (fill[1][2] - fill[0][2]) * ratio)
                    draw.line([(x, y_offset + dy), (x + BOX_W, y_offset + dy)], fill=(r, g, b))
                mask = Image.new("L", (img_w, img_h), 255)
                mask_draw = ImageDraw.Draw(mask)
                mask_draw.rounded_rectangle((0, 0, img_w, img_h), radius=10, fill=255)
                mask_draw.rectangle((0, 0, img_w, img_h), fill=0)
                mask_draw.rounded_rectangle(xy, radius=10, fill=255)
                bg = Image.new("RGB", (img_w, img_h), (0xF8, 0xF9, 0xFA))
                img = Image.composite(img, bg, mask)
                draw = ImageDraw.Draw(img)
                draw.rounded_rectangle(xy, radius=10, fill=None, outline=outline, width=2 if outline else 0)
            else:
                draw.rounded_rectangle(xy, radius=10, fill=fill, outline=outline, width=2 if outline else 0)

            cx = x + BOX_W / 2
            name_text = box.get("name", "")
            tw_name = draw.textlength(name_text, font=font_name)
            draw.text((cx - tw_name / 2, y_offset + 18), name_text, fill=text_clr, font=font_name)

            sub_text = box.get("subtitle", "")
            if sub_text:
                tw_sub = draw.textlength(sub_text, font=font_sub)
                sub_clr = tuple(int(c * 0.9) for c in text_clr) if text_clr == (255, 255, 255) else text_clr
                draw.text((cx - tw_sub / 2, y_offset + 42), sub_text, fill=sub_clr, font=font_sub)

            det_text = box.get("details", "")
            if det_text:
                tw_det = draw.textlength(det_text, font=font_detail)
                det_clr = tuple(int(c * 0.8) for c in text_clr) if text_clr == (255, 255, 255) else text_clr
                draw.text((cx - tw_det / 2, y_offset + 62), det_text, fill=det_clr, font=font_detail)

            positions.append((x, y_offset, x + BOX_W, y_offset + BOX_H))

            if row_idx == sorted_rows[0] and i < len(row_boxes) - 1:
                arrow_x = x + BOX_W + 10
                arrow_mid = y_offset + BOX_H / 2

                n_conn = len(connections)
                conn_spacing = 42 if n_conn <= 2 else max(30, 80 // n_conn)
                conn_block = (n_conn - 1) * conn_spacing
                for ci, conn in enumerate(connections):
                    direction = conn.get("direction", "").upper()
                    protocol = conn.get("protocol", "")
                    ay = arrow_mid - conn_block / 2 + ci * conn_spacing
                    if "READ" in direction:
                        draw.line([(arrow_x + ARROW_W - 10, ay), (arrow_x + 10, ay)],
                                  fill=(0x2E, 0xCC, 0x71), width=2)
                        draw.polygon([(arrow_x + 10, ay - 5), (arrow_x, ay), (arrow_x + 10, ay + 5)],
                                     fill=(0x2E, 0xCC, 0x71))
                        lbl = f"\u2190 {direction}"
                        tw_lbl = draw.textlength(lbl, font=font_arrow)
                        draw.text(((arrow_x + ARROW_W / 2) - tw_lbl / 2, ay - 16),
                                  lbl, fill=(0x2E, 0xCC, 0x71), font=font_arrow)
                    else:
                        draw.line([(arrow_x + 10, ay), (arrow_x + ARROW_W - 10, ay)],
                                  fill=(0xE7, 0x4C, 0x3C), width=2)
                        draw.polygon([(arrow_x + ARROW_W - 10, ay - 5),
                                      (arrow_x + ARROW_W, ay),
                                      (arrow_x + ARROW_W - 10, ay + 5)],
                                     fill=(0xE7, 0x4C, 0x3C))
                        lbl = f"{direction} \u2192"
                        tw_lbl = draw.textlength(lbl, font=font_arrow)
                        draw.text(((arrow_x + ARROW_W / 2) - tw_lbl / 2, ay - 16),
                                  lbl, fill=(0xE7, 0x4C, 0x3C), font=font_arrow)

                    if protocol:
                        tw_p = draw.textlength(f"({protocol})", font=font_proto)
                        draw.text(((arrow_x + ARROW_W / 2) - tw_p / 2, ay + 5),
                                  f"({protocol})", fill=(0x99, 0x99, 0x99), font=font_proto)

                x += BOX_W + GAP_X + ARROW_W
            else:
                x += BOX_W + GAP_X

        row_box_positions[row_idx] = positions
        y_offset += BOX_H + GAP_Y

    buf = BytesIO()
    img.save(buf, format="PNG", dpi=(150, 150))
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Header & Metadata
# ---------------------------------------------------------------------------

def _build_header(doc, meta):
    customer_name = meta.get("customer_name", "[CUSTOMER]")

    p_title = doc.add_paragraph()
    p_title.space_after = Pt(8)
    run = p_title.add_run(f"Quote for {customer_name}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = FONT_BODY
    run.font.color.rgb = CLR_PRIMARY

    meta_fields = [
        ("Quote-Rev", meta.get("quote_rev", "[TBD]-0")),
        ("Expiration Date", meta.get("expiration_date", "[TBD]")),
        ("Contact for Quote", meta.get("contact", "[TBD]")),
        ("Call Reference", meta.get("call_reference", "[TBD]")),
    ]

    table = doc.add_table(rows=2, cols=len(meta_fields))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)

    for i, (label, value) in enumerate(meta_fields):
        label_cell = table.rows[0].cells[i]
        p = label_cell.paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.font.name = FONT_BODY
        run.font.color.rgb = CLR_MUTED
        run.italic = True

        value_cell = table.rows[1].cells[i]
        _set_cell_shading(value_cell, "F8F8F8")
        p = value_cell.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.name = FONT_BODY
        run.font.color.rgb = CLR_BODY

    project_title = meta.get("project_title", "[Project Title]")
    _add_field_label(doc, "Title")
    p_pt = doc.add_paragraph()
    p_pt.space_after = Pt(10)
    run = p_pt.add_run(project_title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = FONT_BODY
    run.font.color.rgb = CLR_BODY


# ---------------------------------------------------------------------------
# Main document assembly
# ---------------------------------------------------------------------------

def generate(json_path, output_path):
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style.font.color.rgb = CLR_BODY
    style.paragraph_format.space_after = Pt(4)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    meta = data.get("meta", {})
    _build_header(doc, meta)

    _build_section1(doc, data)
    _add_separator(doc)
    _build_section2(doc, data)
    _add_separator(doc)
    _build_section3(doc, data)

    _build_quote_details(doc, data)

    doc.add_paragraph()  # spacing before disclaimers
    _build_disclaimers(doc)
    _build_signature(doc)

    doc.save(output_path)
    print(f"Generated: {output_path}")


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python generate_quote_docx.py <content.json> <output.docx>")
        sys.exit(1)
    generate(sys.argv[1], sys.argv[2])
